from __future__ import annotations

from pathlib import Path

from docx import Document
from loguru import logger

from .loader_registry import LoaderRegistry, DocumentContent


@LoaderRegistry.register([".docx"])
class DocxLoader:
    def load(self, file_path: Path) -> DocumentContent:
        doc = Document(str(file_path))
        parts: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # Extract tables
        for table in doc.tables:
            rows: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                parts.append("\n".join(rows))

        full_text = "\n\n".join(parts)
        logger.info(f"DOCX loaded: {file_path.name}, {len(full_text)} chars")
        return DocumentContent(text=full_text)
